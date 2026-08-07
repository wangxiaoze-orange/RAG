"""文档解析模块：MinerU（magic-pdf）清洗优先，失败自动降级 pypdf/pdfplumber/python-docx/纯文本
- prefer：指定解析器（auto/mineru/pypdf/pdfplumber），非 auto 时强制使用（不可用则降级）
- min_confidence：MinerU 块置信度下限，低于阈值的块被过滤；返回平均置信度落库
返回 (markdown_text, parser_name, avg_confidence|None)
"""
import io
import logging

logger = logging.getLogger(__name__)


def _mineru_available() -> bool:
    try:
        import magic_pdf  # noqa: F401
        return True
    except ImportError:
        return False


def _parse_with_mineru(pdf_bytes: bytes, min_confidence: float = 0.0) -> tuple[str, float | None]:
    """MinerU 解析 PDF → Markdown（含表格/标题/行级信息），并按块置信度过滤
    返回 (text, avg_confidence)；置信度字段缺失时 avg_confidence=None
    注意：magic-pdf 各版本 API 有差异，此处做了防御式调用；失败由调用方降级
    """
    from magic_pdf.data.data_reader_writer import BytesDataReader
    from magic_pdf.data.dataset import PymuDocDataset

    ds = PymuDocDataset(BytesDataReader(pdf_bytes), is_debug=False)
    is_text = ds.classify() == "text"
    if is_text:
        pipe = ds.pipe_txt_mode()
    else:
        pipe = ds.pipe_ocr_mode()

    # 置信度过滤：优先用 content_list（各版本字段名不一，防御式探测）
    try:
        content_list = pipe.get_content_list()
        scores: list[float] = []
        kept_parts: list[str] = []
        has_score = False
        for block in content_list:
            score = block.get("score") if isinstance(block, dict) else None
            text_part = block.get("text") if isinstance(block, dict) else str(block)
            if score is not None:
                has_score = True
                scores.append(float(score))
                if float(score) < min_confidence:
                    continue  # 低置信度块丢弃（扫描件噪点/OCR 乱码）
            if text_part and str(text_part).strip():
                kept_parts.append(str(text_part).strip())
        if has_score and kept_parts:
            avg = round(sum(scores) / len(scores), 4) if scores else None
            logger.info("MinerU 置信度过滤: 平均=%.3f 下限=%.2f", avg or 0, min_confidence)
            return "\n\n".join(kept_parts), avg
    except Exception as e:  # noqa: BLE001
        logger.warning("MinerU content_list 不可用，回退 get_markdown: %s", e)
    return pipe.get_markdown(), None


def _parse_pdf_pypdf(pdf_bytes: bytes) -> str:
    from pypdf import PdfReader
    reader = PdfReader(io.BytesIO(pdf_bytes))
    pages = [(page.extract_text() or "") for page in reader.pages]
    return "\n\n".join(pages)


def _parse_pdf_pdfplumber(pdf_bytes: bytes) -> str:
    import pdfplumber
    with pdfplumber.open(io.BytesIO(pdf_bytes)) as pdf:
        pages = [(page.extract_text() or "") for page in pdf.pages]
    return "\n\n".join(pages)


def _parse_docx(data: bytes) -> str:
    import docx
    document = docx.Document(io.BytesIO(data))
    parts = [p.text for p in document.paragraphs if p.text.strip()]
    for table in document.tables:
        for row in table.rows:
            parts.append(" | ".join(cell.text.strip() for cell in row.cells))
    return "\n".join(parts)


def _decode_text(data: bytes) -> str:
    """兼容 UTF-8 / GBK / GB18030 编码"""
    for enc in ("utf-8", "gb18030"):
        try:
            return data.decode(enc)
        except UnicodeDecodeError:
            continue
    return data.decode("utf-8", errors="replace")


def parse_document(
    filename: str,
    data: bytes,
    prefer: str = "auto",
    min_confidence: float = 0.0,
) -> tuple[str, str, float | None]:
    """解析文档为 Markdown 文本，返回 (text, parser_name, avg_confidence)
    prefer: auto（MinerU 优先降级链）/ mineru / pypdf / pdfplumber（强制指定，失败才降级）
    抛出 ValueError：不支持的格式
    """
    ext = filename.rsplit(".", 1)[-1].lower() if "." in filename else ""
    text: str | None = None
    parser = ""
    confidence: float | None = None

    if ext == "pdf":
        # 强制指定降级解析器
        if prefer == "pypdf":
            try:
                text = _parse_pdf_pypdf(data)
                parser = "pypdf"
            except Exception as e:  # noqa: BLE001
                logger.warning("指定 pypdf 失败，降级 pdfplumber: %s", e)
        elif prefer == "pdfplumber":
            try:
                text = _parse_pdf_pdfplumber(data)
                parser = "pdfplumber"
            except Exception as e:  # noqa: BLE001
                logger.warning("指定 pdfplumber 失败，降级 pypdf: %s", e)
        else:
            # auto / mineru：MinerU 优先（安装了 magic-pdf 才尝试）
            if prefer == "mineru" or _mineru_available():
                try:
                    text, confidence = _parse_with_mineru(data, min_confidence)
                    parser = "mineru"
                    logger.info("MinerU 解析成功: %s", filename)
                except Exception as e:  # noqa: BLE001
                    logger.warning("MinerU 解析失败，降级: %s", e)
                    if prefer == "mineru":
                        raise ValueError(f"指定 MinerU 解析失败: {e}") from e
        if text is None:
            try:
                text = _parse_pdf_pypdf(data)
                parser = "pypdf"
            except Exception as e:  # noqa: BLE001
                logger.warning("pypdf 解析失败，降级 pdfplumber: %s", e)
                text = _parse_pdf_pdfplumber(data)
                parser = "pdfplumber"
    elif ext in ("docx", "doc"):
        try:
            text = _parse_docx(data)
            parser = "docx"
        except Exception as e:  # noqa: BLE001
            logger.warning("docx 解析失败，按文本处理: %s", e)
    elif ext in ("txt", "md", "markdown", "text"):
        text = _decode_text(data)
        parser = "text"
    else:
        # 其余（png/jpg 等图片）尝试 OCR 需要 MinerU，v1 暂不支持 → 抛错
        raise ValueError(f"暂不支持的文件类型: {ext or filename}")

    if not text or not text.strip():
        raise ValueError(f"解析结果为空: {filename}（可能是扫描件，建议安装 MinerU 后重试；或调低解析置信度阈值）")
    return text, parser, confidence
